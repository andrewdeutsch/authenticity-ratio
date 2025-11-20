"""
Document processor for brand guidelines uploads.
Supports PDF, DOCX, and TXT formats.
Stores guidelines in AWS S3 with local caching for performance.
"""
import os
import json
import logging
from typing import Optional, Dict, List, Any
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)

# Local cache directory for brand guidelines
LOCAL_CACHE_DIR = Path(__file__).parent.parent / "config" / "brand_guidelines"

# S3 configuration
S3_BUCKET = "ar-ingestion-normalized"
S3_PREFIX = "brand_guidelines"


class BrandGuidelinesProcessor:
    """Process and manage brand guidelines documents with S3 storage"""
    
    def __init__(self, use_s3: bool = True):
        """
        Initialize processor with S3 support.
        
        Args:
            use_s3: If True, use S3 for persistent storage. If False, use local only.
        """
        # Ensure local cache directory exists
        LOCAL_CACHE_DIR.mkdir(parents=True, exist_ok=True)
        logger.info(f"Local cache directory: {LOCAL_CACHE_DIR}")
        
        self.use_s3 = use_s3
        self.s3_client = None
        
        if self.use_s3:
            try:
                import boto3
                from config.database import AWS_CONFIG
                
                self.s3_client = boto3.client(
                    's3',
                    region_name=AWS_CONFIG['region'],
                    aws_access_key_id=AWS_CONFIG['access_key_id'],
                    aws_secret_access_key=AWS_CONFIG['secret_access_key']
                )
                logger.info(f"S3 storage enabled: s3://{S3_BUCKET}/{S3_PREFIX}/")
            except Exception as e:
                logger.warning(f"Failed to initialize S3 client: {e}. Falling back to local storage.")
                self.use_s3 = False
                self.s3_client = None
    
    def extract_text(self, file_path: str, file_type: Optional[str] = None) -> str:
        """
        Extract text from uploaded document.
        
        Args:
            file_path: Path to the document file
            file_type: File extension (pdf, docx, txt). Auto-detected if None.
        
        Returns:
            Extracted text content
        """
        if file_type is None:
            file_type = Path(file_path).suffix.lower().lstrip('.')
        
        logger.info(f"Extracting text from {file_type} file: {file_path}")
        
        try:
            if file_type == 'txt':
                return self._extract_from_txt(file_path)
            elif file_type == 'pdf':
                return self._extract_from_pdf(file_path)
            elif file_type in ['docx', 'doc']:
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            raise
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file with encoding detection"""
        try:
            # Try UTF-8 first
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback to latin-1
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                logger.error(f"Failed to read TXT file with multiple encodings: {e}")
                raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            import PyPDF2
        except ImportError:
            raise ImportError(
                "PyPDF2 is required for PDF processing. "
                "Install with: pip install PyPDF2"
            )
        
        text_parts = []
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                num_pages = len(pdf_reader.pages)
                logger.info(f"Processing PDF with {num_pages} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {e}")
                        continue
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from PDF")
            return full_text
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install with: pip install python-docx"
            )
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            return full_text
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise
    
    def chunk_text(self, text: str, max_chunk_size: int = 2000, overlap: int = 200) -> List[str]:
        """
        Split large documents into manageable chunks with overlap.
        
        Args:
            text: Full text to chunk
            max_chunk_size: Maximum characters per chunk
            overlap: Number of characters to overlap between chunks
        
        Returns:
            List of text chunks
        """
        if len(text) <= max_chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                sentence_ends = ['.', '!', '?', '\n']
                for i in range(end, max(start, end - 200), -1):
                    if text[i] in sentence_ends:
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - overlap if end < len(text) else end
        
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks
    
    def save_guidelines(
        self, 
        brand_id: str, 
        text: str, 
        original_filename: str,
        file_size: int
    ) -> Dict[str, Any]:
        """
        Save processed guidelines to S3 and local cache.
        
        Args:
            brand_id: Brand identifier
            text: Processed text content
            original_filename: Original uploaded filename
            file_size: Size of original file in bytes
        
        Returns:
            Metadata dictionary
        """
        # Create metadata
        word_count = len(text.split())
        metadata = {
            "brand_id": brand_id,
            "original_filename": original_filename,
            "upload_date": datetime.now().isoformat(),
            "file_size_bytes": file_size,
            "word_count": word_count,
            "character_count": len(text),
            "version": "1.0",
            "storage": "s3" if self.use_s3 and self.s3_client else "local"
        }
        
        # Save to S3 if enabled
        if self.use_s3 and self.s3_client:
            try:
                # Upload text file
                text_key = f"{S3_PREFIX}/{brand_id}/guidelines.txt"
                self.s3_client.put_object(
                    Bucket=S3_BUCKET,
                    Key=text_key,
                    Body=text.encode('utf-8'),
                    ContentType='text/plain',
                    Metadata={
                        'brand_id': brand_id,
                        'upload_date': metadata['upload_date']
                    }
                )
                
                # Upload metadata
                metadata_key = f"{S3_PREFIX}/{brand_id}/metadata.json"
                self.s3_client.put_object(
                    Bucket=S3_BUCKET,
                    Key=metadata_key,
                    Body=json.dumps(metadata, indent=2).encode('utf-8'),
                    ContentType='application/json'
                )
                
                logger.info(f"Saved guidelines to S3: s3://{S3_BUCKET}/{text_key}")
            except Exception as e:
                logger.error(f"Failed to save to S3: {e}. Saving locally only.")
                self.use_s3 = False
        
        # Always save to local cache for performance
        brand_dir = LOCAL_CACHE_DIR / brand_id
        brand_dir.mkdir(parents=True, exist_ok=True)
        
        # Save text
        text_file = brand_dir / "guidelines.txt"
        with open(text_file, 'w', encoding='utf-8') as f:
            f.write(text)
        
        # Save metadata
        metadata_file = brand_dir / "metadata.json"
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2)
        
        logger.info(f"Saved guidelines for {brand_id}: {word_count} words")
        return metadata
    
    def load_guidelines(self, brand_id: str) -> Optional[str]:
        """
        Load guidelines for a brand from S3 or local cache.
        
        Args:
            brand_id: Brand identifier
        
        Returns:
            Guidelines text or None if not found
        """
        # Try local cache first for performance
        text_file = LOCAL_CACHE_DIR / brand_id / "guidelines.txt"
        
        if text_file.exists():
            try:
                with open(text_file, 'r', encoding='utf-8') as f:
                    text = f.read()
                logger.info(f"Loaded guidelines from cache for {brand_id}: {len(text)} characters")
                return text
            except Exception as e:
                logger.warning(f"Failed to read from cache: {e}")
        
        # If not in cache and S3 is enabled, try S3
        if self.use_s3 and self.s3_client:
            try:
                text_key = f"{S3_PREFIX}/{brand_id}/guidelines.txt"
                response = self.s3_client.get_object(
                    Bucket=S3_BUCKET,
                    Key=text_key
                )
                text = response['Body'].read().decode('utf-8')
                
                # Save to local cache for next time
                brand_dir = LOCAL_CACHE_DIR / brand_id
                brand_dir.mkdir(parents=True, exist_ok=True)
                with open(text_file, 'w', encoding='utf-8') as f:
                    f.write(text)
                
                logger.info(f"Loaded guidelines from S3 for {brand_id}: {len(text)} characters")
                return text
            except self.s3_client.exceptions.NoSuchKey:
                logger.debug(f"No guidelines found in S3 for brand: {brand_id}")
                return None
            except Exception as e:
                logger.error(f"Error loading guidelines from S3 for {brand_id}: {e}")
                return None
        
        logger.debug(f"No guidelines found for brand: {brand_id}")
        return None
    
    def load_metadata(self, brand_id: str) -> Optional[Dict[str, Any]]:
        """
        Load metadata for brand guidelines from S3 or local cache.
        
        Args:
            brand_id: Brand identifier
        
        Returns:
            Metadata dictionary or None if not found
        """
        # Try local cache first
        metadata_file = LOCAL_CACHE_DIR / brand_id / "metadata.json"
        
        if metadata_file.exists():
            try:
                with open(metadata_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Failed to read metadata from cache: {e}")
        
        # Try S3 if enabled
        if self.use_s3 and self.s3_client:
            try:
                metadata_key = f"{S3_PREFIX}/{brand_id}/metadata.json"
                response = self.s3_client.get_object(
                    Bucket=S3_BUCKET,
                    Key=metadata_key
                )
                metadata = json.loads(response['Body'].read().decode('utf-8'))
                
                # Save to local cache
                brand_dir = LOCAL_CACHE_DIR / brand_id
                brand_dir.mkdir(parents=True, exist_ok=True)
                with open(metadata_file, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, indent=2)
                
                return metadata
            except self.s3_client.exceptions.NoSuchKey:
                return None
            except Exception as e:
                logger.error(f"Error loading metadata from S3 for {brand_id}: {e}")
                return None
        
        return None
    
    def delete_guidelines(self, brand_id: str) -> bool:
        """
        Delete guidelines for a brand from S3 and local cache.
        
        Args:
            brand_id: Brand identifier
        
        Returns:
            True if deleted, False if not found
        """
        deleted = False
        
        # Delete from S3 if enabled
        if self.use_s3 and self.s3_client:
            try:
                # Delete text file
                text_key = f"{S3_PREFIX}/{brand_id}/guidelines.txt"
                self.s3_client.delete_object(Bucket=S3_BUCKET, Key=text_key)
                
                # Delete metadata
                metadata_key = f"{S3_PREFIX}/{brand_id}/metadata.json"
                self.s3_client.delete_object(Bucket=S3_BUCKET, Key=metadata_key)
                
                logger.info(f"Deleted guidelines from S3 for {brand_id}")
                deleted = True
            except Exception as e:
                logger.error(f"Error deleting from S3 for {brand_id}: {e}")
        
        # Delete from local cache
        brand_dir = LOCAL_CACHE_DIR / brand_id
        if brand_dir.exists():
            try:
                import shutil
                shutil.rmtree(brand_dir)
                logger.info(f"Deleted local cache for {brand_id}")
                deleted = True
            except Exception as e:
                logger.error(f"Error deleting local cache for {brand_id}: {e}")
        
        return deleted
    
    def list_brands_with_guidelines(self) -> List[str]:
        """
        List all brands that have uploaded guidelines (from S3 and local cache).
        
        Returns:
            List of brand IDs
        """
        brands = set()
        
        # Check S3 if enabled
        if self.use_s3 and self.s3_client:
            try:
                response = self.s3_client.list_objects_v2(
                    Bucket=S3_BUCKET,
                    Prefix=f"{S3_PREFIX}/",
                    Delimiter='/'
                )
                
                # Extract brand IDs from common prefixes
                if 'CommonPrefixes' in response:
                    for prefix in response['CommonPrefixes']:
                        # Extract brand_id from prefix like "brand_guidelines/mastercard/"
                        parts = prefix['Prefix'].rstrip('/').split('/')
                        if len(parts) >= 2:
                            brands.add(parts[-1])
            except Exception as e:
                logger.error(f"Error listing brands from S3: {e}")
        
        # Check local cache
        if LOCAL_CACHE_DIR.exists():
            for item in LOCAL_CACHE_DIR.iterdir():
                if item.is_dir() and (item / "guidelines.txt").exists():
                    brands.add(item.name)
        
        return sorted(list(brands))
